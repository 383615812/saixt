# -*- coding: utf-8 -*-
"""
2026好题汇编 DOCX → 题库 自动提取导入脚本
1. 提取DOCX文本
2. 解析题目（复用 parse_huikao 逻辑）
3. 质量过滤
4. 导入 SQLite
"""
import os, re, sys, json, sqlite3
from docx import Document

# ---------- 路径 ----------
DB_PATH = r"E:\saixt\server\data\saixt.db"
BASE = r"E:\saixt"

# 科目 → 好题汇编目录名
SUBJECT_DIRS = {
    '历史': '【好题汇编】备战2026年高中历史学业水平合格考真题分类汇编（全国通用）',
    '政治': '【好题汇编】备战2026年高中政治学业水平合格考真题分类汇编（全国通用）',
    '地理': '【好题汇编】备战2026年高中地理学业水平合格考真题分类汇编（全国通用）',
    '英语': '【好题汇编】备战2026年高中英语学业水平合格考真题分类汇编（全国通用）',
    '语文': '【好题汇编】备战2026年高中语文学业水平合格考真题分类汇编（全国通用）',
    '化学': '【好题汇编】备战2026年高中化学学业水平合格考真题分类汇编（全国通用）',
    '物理': '【好题汇编】备战2026年高中物理学业水平合格考真题分类汇编（全国通用）',
    '生物': '【好题汇编】备战2026年高中生物学业水平合格考真题分类汇编（全国通用）',
    '数学': '【好题汇编】备战2026年高中数学学业水平合格考真题分类汇编（全国通用）',
}

# ---------- 正则 ----------
Q_RE = re.compile(r'^\s*(\d{1,3})\s*[．.、]\s*(.*)$')
OPT_RE = re.compile(r'^\s*([A-HＡ-Ｈ])\s*[．.、]\s*(.*)$')
ANS_BLOCK = re.compile(r'(?:【答案】|参考答案[:：])\s*(.*)$')
ANA_BLOCK = re.compile(r'【(解析|详解|分析|知识点|点睛|导语|考点|命题意图|易错警示|名师点睛)】\s*(.*)$')
ANA_START_RE = re.compile(r'^\s*\d{1,3}\s*[．.、]\s*(?:本*)(?:本题考查|本题|此题|该题|这道题|这道|本小题|本段|本句|考查|.{0,8}题[。，,\.])')
ANA_SUB_RE = re.compile(r'^\s*\d{1,3}\s*[．.、]\s*(词汇积累|句式拓展|句型拓展|段落续写|续写线索|词汇激活|行为类|情绪类|高分句型|同义句|原句|拓展句)')
SRC_RE = re.compile(r'[（(]\s*((?:20\d{2}|\d{2}-\d{2})[^）)]{1,30})[）)]')
SECTION_RE = re.compile(r'^(专题|考点|一、|二、|三、|四、|五、|六、|七、|八、|九、|十、|第[一二三四五六七八九十]+部分|（一）|（二）|（三）|（四）|（五）|A卷|B卷|C卷|【精选|直接默写|理解性默写|一\.|二\.|三\.|四\.|五\.)')
INLINE_OPT_RE = re.compile(r'([A-HＡ-Ｈ])\s*[．.、]\s*')
ESSAY_PROMPT_RE = re.compile(r'假定你是|假如你是|假设你是|写一封|写一篇|写一封信|写一则|回信|回一封|书信|写信|征文|倡议书|演讲稿|发言稿|通知|报道|日记|便条|留言条|邀请信|建议信|申请信|感谢信|道歉信|求助信|介绍信|祝贺信|慰问信|推荐信|咨询信|投诉信|启事|海报')

def clean(s):
    s = s.replace('\u3000', ' ').replace('\xa0', ' ')
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def extract_source(text):
    m = SRC_RE.search(text)
    return m.group(1).strip() if m else ''

def split_num_content(text, keep_leading=False):
    text = text.strip()
    if not text:
        return {}
    parts = re.split(r'(?:^|\s)(\d{1,3})\s*[．.、]\s*', text)
    if len(parts) < 3:
        return {None: text}
    result = {}
    if keep_leading and parts[0].strip():
        result[None] = parts[0].strip()
    i = 1
    while i < len(parts) - 1:
        num = int(parts[i])
        content = parts[i + 1].strip()
        if num in result:
            result[num] += '\n' + content
        else:
            result[num] = content
        i += 2
    return result

def split_inline_options(text):
    matches = list(INLINE_OPT_RE.finditer(text))
    if len(matches) < 2:
        return []
    opts = []
    for idx, m in enumerate(matches):
        start = m.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        content = text[start:end].strip()
        if content:
            opts.append((m.group(1).upper(), content))
    return opts

# ---------- 解析器 ----------
class Parser:
    def __init__(self, subject, chapter):
        self.subject = subject
        self.chapter = chapter
        self.pending_q = []
        self.pending_material = None
        self.material_q_count = 0
        self.cur = None
        self.state = 'IDLE'
        self.ans_buf = ''
        self.ana_buf = ''
        self.zone_start = 0
        self.zone_end = 0

    def flush_material(self):
        if self.pending_material and self.material_q_count == 0:
            num = max([q['num'] for q in self.pending_q], default=0) + 1
            self.cur = {'num': num, 'stem': clean(self.pending_material), 'options': [], 'answer': '', 'analysis': '', 'source': extract_source(self.pending_material)}
            self.pending_q.append(self.cur)
            self.state = 'QUESTION'
            self.material_q_count += 1

    def begin_ans_zone(self):
        self.zone_start = self.zone_end
        self.zone_end = len(self.pending_q)

    def assign_to_pending(self, text, field):
        text = text.strip()
        if not text:
            return
        zone = self.pending_q[self.zone_start:self.zone_end]
        if len(zone) == 1 and field == 'analysis':
            q = zone[0]
            if not q[field]:
                q[field] = text
            elif text not in q[field]:
                q[field] += '\n' + text
            return
        mapping = split_num_content(text) if field == 'answer' else split_num_content(text, keep_leading=True)
        if not mapping:
            return
        if None in mapping:
            nums = [n for n in mapping if n is not None]
            target = min(nums) if nums else None
            for q in zone:
                if target is not None:
                    if q['num'] == target and not q[field]:
                        q[field] = mapping[None]
                        break
                elif not q[field]:
                    q[field] = mapping[None]
                    break
        for q in zone:
            if q['num'] in mapping:
                new = mapping[q['num']]
                if not q[field]:
                    q[field] = new
                elif field == 'analysis' and new not in q[field]:
                    q[field] += '\n' + new

    def settle(self):
        if self.ans_buf.strip():
            self.assign_to_pending(self.ans_buf, 'answer')
            self.ans_buf = ''
        if self.ana_buf.strip():
            self.assign_to_pending(self.ana_buf, 'analysis')
            self.ana_buf = ''

    def finalize_cur(self):
        if not self.cur:
            return
        q = self.cur
        q['stem'] = clean(q['stem'])
        seen = set()
        opts = []
        for letter, text in q['options']:
            if letter in seen:
                continue
            seen.add(letter)
            opts.append([letter, clean(text)])
        q['options'] = opts
        self.cur = None

    def handle_line(self, raw):
        s = raw.rstrip('\n').strip()
        if not s:
            return
        if SECTION_RE.match(s) and len(s) < 60:
            self.settle()
            self.finalize_cur()
            self.flush_material()
            self.pending_material = None
            self.material_q_count = 0
            self.state = 'IDLE'
            return
        ma = ANS_BLOCK.match(s)
        if ma:
            self.settle()
            self.finalize_cur()
            self.flush_material()
            self.pending_material = None
            self.material_q_count = 0
            self.begin_ans_zone()
            self.ans_buf = ma.group(1)
            self.state = 'ANSWER'
            return
        mb = ANA_BLOCK.match(s)
        if mb:
            self.settle()
            self.finalize_cur()
            self.ana_buf = mb.group(2)
            self.state = 'ANALYSIS'
            return
        if re.match(r'^[（(]\s*20\d{2}', s) and not Q_RE.match(s):
            self.settle()
            self.finalize_cur()
            self.pending_material = s
            self.material_q_count = 0
            self.state = 'IDLE'
            return
        if self.state in ('ANSWER', 'ANALYSIS'):
            mq = Q_RE.match(s)
            if mq and (SRC_RE.search(s) or self.pending_material):
                pass
            elif mq and (ANA_START_RE.match(s) or ANA_SUB_RE.match(s)):
                self.ana_buf += ' ' + s
                self.state = 'ANALYSIS'
                return
            elif mq:
                self.settle()
                self.ans_buf = s
                self.state = 'ANSWER'
                return
            else:
                if self.state == 'ANSWER':
                    self.ans_buf += ' ' + s
                else:
                    self.ana_buf += ' ' + s
                return
        mq = Q_RE.match(s)
        if mq:
            if not SRC_RE.search(s) and not self.pending_material:
                self.settle()
                self.finalize_cur()
                self.begin_ans_zone()
                self.ans_buf = s
                self.state = 'ANSWER'
                return
            if self.pending_material and (ESSAY_PROMPT_RE.search(self.pending_material) or '注意' in self.pending_material) and not SRC_RE.search(s) and not split_inline_options(mq.group(2)):
                self.pending_material += ' ' + s
                return
            num = int(mq.group(1))
            rest = mq.group(2)
            self.settle()
            self.finalize_cur()
            opts = split_inline_options(rest)
            stem = re.sub(r'^\s*[A-HＡ-Ｈ]\s*[．.、]\s*.*$', '', rest).strip()
            if opts:
                stem = ''
            src = extract_source(s)
            if self.pending_material:
                stem = (self.pending_material + '\n' + stem).strip() if stem else self.pending_material
                if not src:
                    src = extract_source(self.pending_material)
                self.material_q_count += 1
            self.cur = {'num': num, 'stem': stem, 'options': [], 'answer': '', 'analysis': '', 'source': src}
            for letter, text in opts:
                self.cur['options'].append([letter, text])
            self.pending_q.append(self.cur)
            self.state = 'QUESTION'
            return
        mo = OPT_RE.match(s)
        if mo and self.state == 'QUESTION':
            letter = mo.group(1).upper()
            content = mo.group(2)
            inline = split_inline_options(s)
            if inline:
                for l, t in inline:
                    self.cur['options'].append([l, t])
            else:
                self.cur['options'].append([letter, content])
            return
        if self.state == 'QUESTION':
            self.cur['stem'] += ' ' + s
        elif self.state == 'ANSWER':
            self.ans_buf += ' ' + s
        elif self.state == 'ANALYSIS':
            self.ana_buf += ' ' + s
        else:
            if self.pending_material:
                self.pending_material += ' ' + s

    def finish(self):
        self.settle()
        self.finalize_cur()
        for q in self.pending_q:
            if not q['answer']:
                q['answer'] = '未提供'
            if not q['analysis']:
                q['analysis'] = ''
        return self.pending_q

def extract_docx_text(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace('\n', ' / ') for c in row.cells]
            line = " | ".join(cells)
            if line.strip(" |"):
                lines.append(line)
    return lines

def parse_docx(path, subject, chapter):
    lines = extract_docx_text(path)
    p = Parser(subject, chapter)
    for line in lines:
        p.handle_line(line)
    return p.finish()

def get_chapter_name(filename):
    """从文件名提取章节名"""
    name = os.path.splitext(os.path.basename(filename))[0]
    name = re.sub(r'（学考真题汇编）', '', name)
    name = re.sub(r'（解析版）', '', name)
    name = re.sub(r'（原卷版）', '', name)
    return name.strip()

def quality_check(q):
    """质量检查：返回 True 表示通过"""
    stem = q.get('stem', '').strip()
    if not stem or len(stem) < 5:
        return False
    # 客观题必须有选项（≥2个），主观题可以无选项
    opts = q.get('options', [])
    if opts:
        if len(opts) < 2:
            return False
        # 检查空选项（公式丢失导致）
        empty_opts = sum(1 for o in opts if not o[1].strip())
        if empty_opts > 0:
            return False
        # 单选项（只有A或只有A和B且内容为空）
        if len(opts) == 1:
            return False
    answer = q.get('answer', '').strip()
    if not answer or answer == '未提供':
        return False
    analysis = q.get('analysis', '').strip()
    if not analysis:
        return False
    return True

def qtype_of(q):
    opts = q.get('options', [])
    if not opts:
        return 'subjective'
    answer = q.get('answer', '').strip().upper()
    # 多选题：答案含多个字母
    letters = [c for c in answer if c in 'ABCDEFGH']
    if len(letters) >= 2 and len(answer) <= len(''.join(letters)):
        return 'multiple'
    # 判断题：选项只有A/B且答案为A或B
    if len(opts) == 2 and answer in ('A', 'B'):
        return 'judge'
    return 'single'

def main():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    total_imported = 0
    total_skipped = 0
    subject_stats = {}

    for subject, dirname in SUBJECT_DIRS.items():
        dirpath = os.path.join(BASE, dirname)
        if not os.path.isdir(dirpath):
            continue
        # 只处理解析版（含答案）
        docx_files = []
        for root, dirs, files in os.walk(dirpath):
            for f in files:
                if f.endswith('.docx') and '解析' in f:
                    docx_files.append(os.path.join(root, f))

        subj_total = 0
        subj_skip = 0
        for docx_path in sorted(docx_files):
            chapter = get_chapter_name(docx_path)
            try:
                questions = parse_docx(docx_path, subject, chapter)
            except Exception as e:
                print(f'[ERR] {subject}/{chapter}: {e}')
                continue

            for q in questions:
                if not quality_check(q):
                    subj_skip += 1
                    total_skipped += 1
                    continue

                qtype = qtype_of(q)
                opts_json = json.dumps(q['options'], ensure_ascii=False)
                source = q.get('source', '') or '2026好题汇编'
                difficulty = 2

                # 去重：检查题干是否已存在
                stem_check = q['stem'][:50].strip()
                if not stem_check:
                    continue
                existing = cur.execute(
                    "SELECT id FROM questions WHERE subject=? AND stem LIKE ? LIMIT 1",
                    (subject, stem_check + '%')
                ).fetchone()
                if existing:
                    subj_skip += 1
                    total_skipped += 1
                    continue

                cur.execute(
                    "INSERT INTO questions (subject, chapter, type, difficulty, stem, options, answer, analysis, source) VALUES (?,?,?,?,?,?,?,?,?)",
                    (subject, chapter, qtype, difficulty, q['stem'], opts_json, q['answer'], q['analysis'], source)
                )
                subj_total += 1
                total_imported += 1

        subject_stats[subject] = {'imported': subj_total, 'skipped': subj_skip}
        print(f'{subject}: 导入 {subj_total} 题, 跳过 {subj_skip} 题')

    conn.commit()

    # 验证总数
    cur.execute("SELECT COUNT(*) FROM questions")
    final_total = cur.fetchone()[0]
    conn.close()

    print(f'\n==== 汇总 ====')
    for s, st in subject_stats.items():
        print(f'{s}: 导入{st["imported"]} 跳过{st["skipped"]}')
    print(f'总导入: {total_imported} 题')
    print(f'总跳过: {total_skipped} 题')
    print(f'题库总数: {final_total} 题')

if __name__ == '__main__':
    main()
