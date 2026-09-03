# -*- coding: utf-8 -*-
"""标准试卷格式解析：题干+内联/独立选项在主文，卷尾"N．答案"+【详解】按题号合并。"""
import re, os
from docx import Document

Q = re.compile(r'^\s*(\d{1,3})\s*[．.、]\s*(.*)$')
OPT_INLINE = re.compile(r'([A-ZＡ-Ｚ])\s*[．.、]\s*')
ANA_LINE = re.compile(r'^\s*【(详解|解析|答案)】\s*(.*)$')
SKIP_PREFIX = ('学校:', '姓名:', '班级:', '考号:', '注意事项', '考生', '第I卷', '第Ⅰ卷',
               '第II卷', '第Ⅱ卷', '一、选', '二、', '三、', '四、', '五、', '六、', '七、',
               '选择题', '非选择题', '单项选择', '多项选择', '判断题', '综合题', '操作题',
               '填空题', '简答题', '解答题', '第一部分', '第二部分')
ANSWER_RE = re.compile(r'^\s*(\d{1,3})\s*[．.、]\s*([A-DＡ-Ｄ])\s*$')


def clean(s):
    s = (s or '').replace('\u3000', ' ').replace('\xa0', ' ').replace('　', ' ')
    return re.sub(r'\s+', ' ', s).strip()


def split_opts(text):
    ms = list(OPT_INLINE.finditer(text))
    if len(ms) < 2:
        return []
    out = []
    for i, m in enumerate(ms):
        start = m.end()
        end = ms[i + 1].start() if i + 1 < len(ms) else len(text)
        c = text[start:end].strip().rstrip('。；,，')
        if c:
            out.append((m.group(1).upper(), c))
    return out


def parse_docx(path):
    doc = Document(path)
    paras = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    for t in doc.tables:
        for row in t.rows:
            j = clean(' '.join(c.text for c in row.cells))
            if j:
                paras.append(j)

    questions = {}    # n -> {'stem','options'}
    order = []
    answers = {}
    analysis = {}
    cur = None
    last_ans = None

    for raw in paras:
        # ---- 答案区：单字母答案行 ----
        ar = ANSWER_RE.match(raw)
        if ar:
            last_ans = int(ar.group(1))
            answers[last_ans] = ar.group(2).upper()
            analysis[last_ans] = ''
            continue
        # 详解/解析行(尾注释)
        ma = ANA_LINE.match(raw)
        if ma and last_ans is not None and raw.startswith('【'):
            c = ma.group(2) or ''
            if c and c != '略':
                analysis[last_ans] = (analysis.get(last_ans) or '') + c
            continue
        # 答案区续行
        if last_ans is not None and not Q.match(raw) and not any(raw.startswith(s) for s in SKIP_PREFIX):
            if raw and raw != '略':
                # 避免吞掉新的题目标题行
                analysis[last_ans] = (analysis.get(last_ans) or '') + raw
            continue

        if any(raw.startswith(s) for s in SKIP_PREFIX):
            continue

        qm = Q.match(raw)
        if qm:
            n = int(qm.group(1))
            content = qm.group(2)
            if len(content) <= 1:
                continue
            if raw and content and len(content) <= 2 and re.match(r'^[A-Da-dＡ-Ｄａ-ｄ]$', content):
                continue
            if n not in questions:
                questions[n] = {'stem': None, 'options': []}
                order.append(n)
                cur = n
            # 题干后续段追加
            if questions[n]['stem'] is None:
                questions[n]['stem'] = content
            else:
                # 该行可能是上一题的行内选项
                inl = split_opts(content)
                if len(inl) >= 2 and not n in answers:
                    questions[n]['options'] += inl
                else:
                    questions[n]['stem'] += ' ' + content
            cur = n
            continue

        # 非题号行：独立选项行，挂到 cur
        inl = split_opts(raw)
        if len(inl) >= 2 and cur is not None:
            questions[cur]['options'] += inl

    # 提取题干内联选项(若题干行自带 A．B．)
    for n in order:
        q = questions[n]
        if q['options']:
            continue
        st = q['stem'] or ''
        inl = split_opts(st)
        if len(inl) >= 2:
            first = OPT_INLINE.search(st)
            if first:
                q['stem'] = st[:first.start()].strip()
            q['options'] = inl

    result = []
    for n in order:
        q = questions[n]
        stem = (q['stem'] or '').strip()
        if not stem or len(stem) < 3:
            continue
        # 归一去重取前4项
        seen, opts = [], []
        for ch, c in q['options']:
            if ch in seen:
                continue
            seen.append(ch)
            opts.append(c)
            if len(seen) >= 4:
                break
        if len(seen) < 2:
            continue
        ans = answers.get(n, '')
        if ans not in 'ABCD':
            continue
        result.append({'num': n, 'stem': stem, 'options': opts,
                       'answer': ans, 'analysis': (analysis.get(n) or '').strip(),
                       'source': paper_source(path)})
    return result


def paper_source(path):
    base = os.path.basename(path).replace('.docx', '')
    base = re.sub(r'\(\d\)$', '', base).strip()
    return base or '真题综合'


def quality_check(q):
    stem = (q.get('stem') or '').strip()
    if not stem or len(stem) < 5:
        return False
    if len(q.get('options') or []) < 2:
        return False
    if not q.get('answer'):
        return False
    return True