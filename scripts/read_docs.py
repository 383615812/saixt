# -*- coding: utf-8 -*-
import sys, glob, os
import docx

def read_docx(path):
    d = docx.Document(path)
    lines = []
    for p in d.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for t in d.tables:
        lines.append('[表格]')
        for row in t.rows:
            lines.append(' | '.join(c.text.strip() for c in row.cells))
    return lines

base = r'E:\saixt'
for f in glob.glob(os.path.join(base, '*.docx')):
    print('=' * 60)
    print('文件:', os.path.basename(f))
    print('=' * 60)
    try:
        for line in read_docx(f):
            print(line)
    except Exception as e:
        print('读取失败:', e)
